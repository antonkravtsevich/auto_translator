import nltk
import pandas as pd
import sys
import requests
from tqdm import tqdm

from docx import Document
from nltk.corpus import stopwords 
from nltk.stem import WordNetLemmatizer
from nltk.tokenize import word_tokenize 
from py_translator import Translator
from get_fresh_proxy import get_fresh_proxy
from json.decoder import JSONDecodeError


class DictionaryGenerator():

    def __init__(self, df, document, words_count, with_translations):
        self.df = df
        self.document = document
        self.words_count = words_count
        self.lemmatizer = WordNetLemmatizer()
        self.dictionary_words = set()
        self.stop_words = set(stopwords.words('english'))
        self.with_translations = with_translations
        self.process_document()
        

    def get_word_position(self, word):
        """
        Get place of word in frequency dictionary
        The more this word is mentioned in english language = the higher the place will be
        """
        try:
            return self.df[self.df['word']==word].index.tolist()[0]    
        except IndexError:
            # word not in dictionary
            return 0


    def get_word_normal_form(self, word):
        word = self.lemmatizer.lemmatize(word)
        word = self.lemmatizer.lemmatize(word, 'v')
        return word


    def update_words_set_by_article(self, article):
        """
        Clean article data and add all new words to words set
        """
        to_replace = '(){}[],.!?"'
        for symbol in to_replace:
            article = article.replace(symbol, ' ')
        article = article.strip()
        article = article.lower()
        article_words = word_tokenize(article)
        
        for word in article_words:
            if '-' in word:
                article_words+=word.split('-')

        filtered_article_words = [word for word in article_words if word not in self.stop_words]
        filtered_article_words = [word for word in filtered_article_words if not word.endswith('lly')]
        normalized_words = [self.get_word_normal_form(word) for word in filtered_article_words]
        self.dictionary_words.update(normalized_words)
    

    def process_document(self):
        table = self.document.tables[0]
        print('Process articles')
        for row in tqdm(table.rows):
            article = row.cells[0].text
            self.update_words_set_by_article(article)


    # TODO: remove this function
    def show_words_debug(self):
        words = list(self.dictionary_words)
        words = sorted(words, key=lambda a: self.get_word_position(a))
        for word in words:
            print('{}: {}'.format(word, self.get_word_position(word)))


    def get_rearest_words(self):
        words = list(self.dictionary_words)
        words = sorted(words, key=lambda a: self.get_word_position(a), reverse=True)[:self.words_count]
        return words

    def get_dictionary(self):
        print('Sort words by frequency')
        words = self.get_rearest_words()
        if not self.with_translations:
            return sorted(words)
        print('Translate words')
        words_with_translation = []
        for word in tqdm(sorted(words)):
            translated=False
            translator = Translator(proxies=get_fresh_proxy())
            print('word: {}'.format(word))
            while not translated:
                try:
                    words_with_translation.append((word, translator.translate(word)))
                    translated = True
                except JSONDecodeError:
                    translator = Translator(proxies=get_fresh_proxy())
                    print('JSON decode error')     
                except IndexError:
                    translator = Translator(proxies=get_fresh_proxy())
                    print('index error')
                except requests.exceptions.ProxyError:
                    translator = Translator(proxies=get_fresh_proxy())
                    print('can\'t connect to proxy')

        return words_with_translation

    def write_document(self):
        dictionary = self.get_dictionary()
        output_document = Document()
        table = output_document.add_table(rows=len(dictionary), cols=2)
        for dict_row, row in tqdm(zip(dictionary, table.rows)):
            if self.with_translations:
                row.cells[0].text = dict_row[0]
                row.cells[1].text = dict_row[1]
            else:
                row.cells[0].text = dict_row
        self.document.save('dictionary.docx')


def main():
    try:
        filename = sys.argv[1]
        words_count = int(sys.argv[2])
    except IndexError:
        print('Use "python3 auto_dictionary_generator.py %filename%" %words_count%')
        sys.exit()

    with_translations = False
    if len(sys.argv)==4:
        if sys.argv[3] =='--with-translations':
            with_translations = True


    nltk.download('stopwords')

    df = pd.read_csv('./data/unigram_freq.csv')
    document = Document(filename)

    dictionary_generator = DictionaryGenerator(df, document, words_count, with_translations)
    dictionary_generator.write_document()


if __name__ == '__main__':
    main()